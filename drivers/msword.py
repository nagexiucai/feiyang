#!/usr/bin/env python
#coding=utf-8
'''
@author: nagexiucai
@license: gpl
'''

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from common.fsoperate import DoesExist

class DocxTemplate(object):
    _style = {        
        'default': {
            'font': {'name': u'宋体', 'size': 20}
            }
        }
    def __init__(self):
        self._style = {
            u'封面': {},
            u'眉脚': {},
            u'正文': {},
            u'插图': {},
            u'表格': {},
            u'目录': {}
            }
    def GetStyle(self, what):
        return self._style.get(what) #DocxTemplate._style.get('default').get(what)

class MSWord(object):
    def __init__(self):
        self.template = DocxTemplate()
    def SetTemplate(self, template):
        self.template = template
    def CreateDocx(self, data, out):
        #入参：data [(层级, 章节标题, 插图JPEG文件绝对路径, 文本),...]
        #入参：out Word文件绝对路径
        #TODO: 丰富排版风格
        document = Document()
        for i, caption, jpeg, text in data:
            document.add_heading(caption, i)
            if jpeg is not None and DoesExist(jpeg):
                document.add_picture(jpeg)
            paragraph = document.add_paragraph()
            run = paragraph.add_run(text)
            run.font.size = Pt(20)
            run.font.name = u'宋体'
        document.save(out)
